"""
DDR Report Generator
Reads Inspection Report + Thermal Report PDFs and generates a structured DDR (.docx)
Works generically with any similar inspection + thermal PDF pair.
"""

import fitz  # pymupdf
import os
import re
import json
import argparse
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google import genai
from google.genai import errors as genai_errors


def load_dotenv_file(dotenv_path: Path) -> None:
    """Load KEY=VALUE pairs from a .env file into os.environ without overriding existing vars."""
    if not dotenv_path.exists() or not dotenv_path.is_file():
        return

    for raw_line in dotenv_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue

        if line.startswith("export "):
            line = line[len("export "):].strip()

        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def auto_discover_pdfs(base_dir: Path) -> tuple[str | None, str | None]:
    """Best-effort PDF discovery when explicit CLI paths are not provided."""
    preferred_pairs = [
        ("Sample_Report.pdf", "Thermal_Images.pdf"),
        ("Inspection_Report.pdf", "Thermal_Report.pdf"),
    ]
    for inspection_name, thermal_name in preferred_pairs:
        inspection_path = base_dir / inspection_name
        thermal_path = base_dir / thermal_name
        if inspection_path.is_file() and thermal_path.is_file():
            return str(inspection_path), str(thermal_path)

    pdfs = sorted(base_dir.glob("*.pdf"))
    if len(pdfs) < 2:
        return None, None

    # Prefer names that hint at report type; fallback to first two PDFs.
    inspection_pdf = next((p for p in pdfs if "thermal" not in p.name.lower()), pdfs[0])
    thermal_pdf = next((p for p in pdfs if "thermal" in p.name.lower()), None)
    if thermal_pdf is None:
        thermal_pdf = pdfs[1] if pdfs[1] != inspection_pdf else (pdfs[2] if len(pdfs) > 2 else None)

    if thermal_pdf is None or inspection_pdf == thermal_pdf:
        return None, None
    return str(inspection_pdf), str(thermal_pdf)


load_dotenv_file(Path(__file__).resolve().parent / ".env")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")  # set via env var
BASE_DIR = Path(__file__).resolve().parent
OUT_DIR = str(BASE_DIR / "output")
IMG_DIR = str(BASE_DIR / "extracted_images")
ENV_INSPECTION_PDF = os.environ.get("INSPECTION_PDF", "").strip()
ENV_THERMAL_PDF = os.environ.get("THERMAL_PDF", "").strip()
ENV_DDR_OUT_PATH = os.environ.get("DDR_OUTPUT_PATH", "").strip()
ENV_GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "").strip()
ENV_ALLOW_FALLBACK = os.environ.get("DDR_ALLOW_FALLBACK", "1").strip().lower() in {"1", "true", "yes", "y", "on"}


# STEP 1: EXTRACT TEXT FROM PDFs
def extract_text(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text_parts = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            text_parts.append(f"[PAGE {page_num+1}]\n{text}")
    return "\n\n".join(text_parts)



# STEP 2: EXTRACT IMAGES

def extract_inspection_images(pdf_path: str, out_dir: str) -> list:
    """Extract embedded site photos from inspection PDF."""
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    saved = []
    seen = set()
    for page_num in range(len(doc)):
        page = doc[page_num]
        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            base = doc.extract_image(xref)
            w, h = base["width"], base["height"]
            aspect = w / h if h > 0 else 99
            if w < 300 or h < 200 or aspect > 5:
                continue
            fname = f"insp_p{page_num+1:02d}_{len(saved)+1:02d}.{base['ext']}"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, "wb") as f:
                f.write(base["image"])
            saved.append({"filename": fname, "page": page_num + 1, "path": fpath})
    return saved


def render_thermal_pages(pdf_path: str, out_dir: str) -> list:
    """Render each thermal PDF page as image to capture full layout."""
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    saved = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        fname = f"thermal_page_{page_num+1:02d}.png"
        fpath = os.path.join(out_dir, fname)
        pix.save(fpath)
        saved.append({"filename": fname, "page": page_num + 1, "path": fpath})
    return saved



# STEP 3: BUILD DDR PROMPT & CALL GEMINI

DDR_PROMPT = """
You are an expert building diagnostic engineer generating a professional DDR (Detailed Diagnostic Report).

You have been given:
1. An Inspection Report (text content below)
2. A Thermal Report (text content below)

Your task: Generate a structured DDR in **valid JSON** format with EXACTLY these 7 sections.
Use ONLY information found in the provided documents. Do NOT invent facts.

OUTPUT FORMAT (respond with JSON only, no markdown fences, no extra text):
{
  "property_issue_summary": "2-3 paragraph plain English summary of the main problems found",
  
  "area_wise_observations": [
    {
      "area": "Area name (e.g., Hall, Master Bedroom, Kitchen)",
      "negative_side": "What was observed on the impacted/interior side",
      "positive_side": "What was observed on the source/exposed side",
      "thermal_finding": "Relevant thermal reading or finding for this area, or 'Not Available'",
      "image_reference": "Which inspection photo numbers relate (e.g., Photo 1-7), or 'Not Available'"
    }
  ],
  
  "probable_root_cause": "Clear explanation of root causes. Merge inspection + thermal data logically. If conflict, state it.",
  
  "severity_assessment": [
    {
      "area": "Area name",
      "severity": "High / Medium / Low",
      "reasoning": "Why this severity rating"
    }
  ],
  
  "recommended_actions": [
    {
      "priority": "Immediate / Short-term / Long-term",
      "action": "Specific action to take",
      "area": "Applicable area"
    }
  ],
  
  "additional_notes": "Any other relevant findings, checklist flags, inspection score, property details",
  
  "missing_or_unclear_information": [
    "Item that is missing or unclear — write 'Not Available' for fields with no data"
  ]
}

--- INSPECTION REPORT TEXT ---
{inspection_text}

--- THERMAL REPORT TEXT ---
{thermal_text}
"""


def parse_paged_text(extracted_text: str) -> dict[int, str]:
    """Parse [PAGE n] blocks into a page->text dictionary."""
    page_map: dict[int, str] = {}
    matches = list(re.finditer(r"\[PAGE\s+(\d+)\]\n", extracted_text))
    for idx, match in enumerate(matches):
        page_num = int(match.group(1))
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(extracted_text)
        page_map[page_num] = extracted_text[start:end].strip()
    return page_map


def normalize_ddr_payload(ddr: dict, inspection_text: str, thermal_text: str) -> dict:
    """Validate and normalize DDR payload for consistent downstream rendering."""
    if not isinstance(ddr, dict):
        ddr = {}

    def norm_text(value: object, default: str = "Not Available") -> str:
        if not isinstance(value, str):
            return default
        cleaned = value.strip()
        return cleaned if cleaned else default

    area_obs_raw = ddr.get("area_wise_observations", [])
    if not isinstance(area_obs_raw, list):
        area_obs_raw = []

    normalized_obs = []
    seen_obs = set()
    for item in area_obs_raw:
        if not isinstance(item, dict):
            continue
        obs = {
            "area": norm_text(item.get("area")),
            "negative_side": norm_text(item.get("negative_side")),
            "positive_side": norm_text(item.get("positive_side")),
            "thermal_finding": norm_text(item.get("thermal_finding")),
            "image_reference": norm_text(item.get("image_reference"), default="Image Not Available"),
        }
        signature = (
            obs["area"].lower(),
            obs["negative_side"].lower(),
            obs["positive_side"].lower(),
            obs["thermal_finding"].lower(),
        )
        if signature in seen_obs:
            continue
        seen_obs.add(signature)
        normalized_obs.append(obs)

    if not normalized_obs:
        normalized_obs = [{
            "area": "General",
            "negative_side": "Not Available",
            "positive_side": "Not Available",
            "thermal_finding": "Not Available",
            "image_reference": "Image Not Available",
        }]

    severity_raw = ddr.get("severity_assessment", [])
    if not isinstance(severity_raw, list):
        severity_raw = []

    normalized_severity = []
    seen_severity = set()
    for item in severity_raw:
        if not isinstance(item, dict):
            continue
        area = norm_text(item.get("area"))
        if area.lower() in seen_severity:
            continue
        seen_severity.add(area.lower())
        sev = norm_text(item.get("severity"), default="Medium")
        if sev not in {"High", "Medium", "Low"}:
            sev = "Medium"
        normalized_severity.append({
            "area": area,
            "severity": sev,
            "reasoning": norm_text(item.get("reasoning")),
        })

    actions_raw = ddr.get("recommended_actions", [])
    if not isinstance(actions_raw, list):
        actions_raw = []

    normalized_actions = []
    seen_actions = set()
    for item in actions_raw:
        if not isinstance(item, dict):
            continue
        action_item = {
            "priority": norm_text(item.get("priority"), default="Short-term"),
            "action": norm_text(item.get("action")),
            "area": norm_text(item.get("area")),
        }
        if action_item["priority"] not in {"Immediate", "Short-term", "Long-term"}:
            action_item["priority"] = "Short-term"
        action_signature = (
            action_item["priority"].lower(),
            action_item["area"].lower(),
            action_item["action"].lower(),
        )
        if action_signature in seen_actions:
            continue
        seen_actions.add(action_signature)
        normalized_actions.append(action_item)

    missing_raw = ddr.get("missing_or_unclear_information", [])
    if not isinstance(missing_raw, list):
        missing_raw = []
    missing_info = [norm_text(item) for item in missing_raw if isinstance(item, str) and item.strip()]

    combined_text = f"{inspection_text} {thermal_text}".lower()
    has_positive_moisture_claim = any(
        key in combined_text for key in ["no leakage", "no seepage", "dry wall", "no damp"]
    )
    has_negative_moisture_claim = any(
        key in combined_text for key in ["leak", "seepage", "damp", "moisture", "water ingress"]
    )
    if has_positive_moisture_claim and has_negative_moisture_claim:
        conflict_note = "Potentially conflicting moisture statements observed across source documents."
        if conflict_note not in missing_info:
            missing_info.append(conflict_note)

    if not missing_info:
        missing_info = ["Not Available"]

    return {
        "property_issue_summary": norm_text(ddr.get("property_issue_summary")),
        "area_wise_observations": normalized_obs,
        "probable_root_cause": norm_text(ddr.get("probable_root_cause")),
        "severity_assessment": normalized_severity,
        "recommended_actions": normalized_actions,
        "additional_notes": norm_text(ddr.get("additional_notes")),
        "missing_or_unclear_information": missing_info,
    }


def build_area_image_maps(
    areas: list,
    inspection_images: list,
    thermal_images: list,
    inspection_page_map: dict[int, str],
    thermal_page_map: dict[int, str],
) -> tuple[dict[int, list], dict[int, dict]]:
    """Map area observations to likely relevant inspection and thermal images."""
    insp_img_map: dict[int, list] = {}
    thermal_img_map: dict[int, dict] = {}

    for idx, obs in enumerate(areas):
        area_name = str(obs.get("area", "")).strip().lower()

        matched_inspection = []
        if area_name and area_name != "not available":
            for img in inspection_images:
                page_text = inspection_page_map.get(img.get("page", -1), "").lower()
                if area_name in page_text:
                    matched_inspection.append(img)

        if not matched_inspection:
            start = idx * 2
            matched_inspection = inspection_images[start:start + 2] if start < len(inspection_images) else []
        insp_img_map[idx] = matched_inspection[:2]

        matched_thermal = None
        if area_name and area_name != "not available":
            for img in thermal_images:
                page_text = thermal_page_map.get(img.get("page", -1), "").lower()
                if area_name in page_text:
                    matched_thermal = img
                    break
        if matched_thermal is None and idx < len(thermal_images):
            matched_thermal = thermal_images[idx]
        if matched_thermal is not None:
            thermal_img_map[idx] = matched_thermal

    return insp_img_map, thermal_img_map


def build_fallback_ddr(inspection_text: str, thermal_text: str, reason: str) -> dict:
    """Create a minimal DDR payload when Gemini is unavailable."""
    area_keywords = [
        "hall", "living", "kitchen", "master bedroom", "bedroom", "bathroom",
        "toilet", "balcony", "passage", "ceiling", "wall", "utility",
    ]
    found_areas = []
    lower_text = f"{inspection_text}\n{thermal_text}".lower()
    for name in area_keywords:
        if name in lower_text:
            title = " ".join(part.capitalize() for part in name.split())
            if title not in found_areas:
                found_areas.append(title)

    if not found_areas:
        found_areas = ["General"]

    observations = []
    severities = []
    actions = []
    for area in found_areas[:8]:
        observations.append({
            "area": area,
            "negative_side": "Detailed inspection-side observation not available in fallback mode.",
            "positive_side": "Detailed source-side observation not available in fallback mode.",
            "thermal_finding": "Not Available",
            "image_reference": "Not Available",
        })
        severities.append({
            "area": area,
            "severity": "Medium",
            "reasoning": "Auto-assigned in fallback mode because Gemini analysis was unavailable.",
        })
        actions.append({
            "priority": "Short-term",
            "action": "Re-run AI analysis after quota is restored and validate area-level details.",
            "area": area,
        })

    inspection_excerpt = inspection_text.strip().replace("\n", " ")[:500]
    thermal_excerpt = thermal_text.strip().replace("\n", " ")[:500]
    summary = (
        "AI analysis could not be completed due to Gemini API unavailability. "
        "A fallback DDR structure has been generated so documentation can proceed.\n\n"
        f"Inspection text length: {len(inspection_text)} chars. Thermal text length: {len(thermal_text)} chars."
    )

    return {
        "property_issue_summary": summary,
        "area_wise_observations": observations,
        "probable_root_cause": (
            "Root cause could not be reliably derived because Gemini analysis failed. "
            "Please rerun once API quota is available."
        ),
        "severity_assessment": severities,
        "recommended_actions": actions,
        "additional_notes": (
            f"Fallback reason: {reason}. "
            f"Inspection excerpt: {inspection_excerpt or 'Not Available'}. "
            f"Thermal excerpt: {thermal_excerpt or 'Not Available'}."
        ),
        "missing_or_unclear_information": [
            "AI-generated detailed findings unavailable due to API quota/rate limit.",
            "Area-level thermal interpretation not available in fallback mode.",
            "Re-run with a working Gemini quota for full analysis.",
        ],
    }


def call_gemini_for_ddr(inspection_text: str, thermal_text: str, api_key: str) -> dict:
    client = genai.Client(api_key=api_key)

    prompt = (
        DDR_PROMPT
        .replace("{inspection_text}", inspection_text[:30000])
        .replace("{thermal_text}", thermal_text[:10000])
    )

    model_candidates = []
    if ENV_GEMINI_MODEL:
        model_candidates.append(ENV_GEMINI_MODEL)
    model_candidates.extend([
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
        "gemini-1.5-flash",
    ])

    # Preserve order while removing duplicates.
    seen = set()
    unique_candidates = []
    for model_name in model_candidates:
        if model_name not in seen:
            seen.add(model_name)
            unique_candidates.append(model_name)

    response = None
    last_error = None
    selected_model = None
    for model_name in unique_candidates:
        try:
            response = client.models.generate_content(model=model_name, contents=prompt)
            selected_model = model_name
            break
        except genai_errors.ClientError as exc:
            last_error = exc
            err_text = str(exc)
            if "NOT_FOUND" in err_text:
                continue
            if "RESOURCE_EXHAUSTED" in err_text or "quota" in err_text.lower():
                raise RuntimeError(
                    "Gemini quota exceeded for this API key/project. "
                    "Update billing/quota or use a different API key/project and retry."
                ) from exc
            raise RuntimeError(f"Gemini API call failed: {err_text}") from exc

    if response is None:
        available_models = []
        try:
            for m in client.models.list():
                name = getattr(m, "name", "")
                if name:
                    available_models.append(name)
                if len(available_models) >= 10:
                    break
        except Exception:
            pass

        available_hint = ""
        if available_models:
            available_hint = f" Available models (first 10): {', '.join(available_models)}"
        raise RuntimeError(
            "No compatible Gemini model found. "
            "Set GEMINI_MODEL in .env to a model available for your API key/project."
            + available_hint
        ) from last_error

    print(f"  Gemini model used: {selected_model}")
    raw = response.text.strip()

    # Strip markdown fences if present
    raw = re.sub(r'^```(?:json)?\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print("Raw response (first 500 chars):", raw[:500])
        raise



# STEP 4: BUILD .DOCX REPORT


SEVERITY_COLORS = {
    "High":   RGBColor(0xC0, 0x00, 0x00),
    "Medium": RGBColor(0xFF, 0x7F, 0x00),
    "Low":    RGBColor(0x00, 0x70, 0x00),
}

PRIORITY_COLORS = {
    "Immediate":   RGBColor(0xC0, 0x00, 0x00),
    "Short-term":  RGBColor(0xFF, 0x7F, 0x00),
    "Long-term":   RGBColor(0x00, 0x70, 0x00),
}


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_section_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_severity_badge(para, severity: str):
    run = para.add_run(f"  [{severity.upper()}]  ")
    run.bold = True
    color = SEVERITY_COLORS.get(severity, RGBColor(0x66, 0x66, 0x66))
    run.font.color.rgb = color


def build_docx_report(
    ddr: dict,
    inspection_images: list,
    thermal_images: list,
    out_path: str,
    inspection_page_map: dict[int, str] | None = None,
    thermal_page_map: dict[int, str] | None = None,
):
    doc = DocxDocument()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover Header ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("DETAILED DIAGNOSTIC REPORT (DDR)")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Property Inspection Report | Prepared by AI DDR System").font.size = Pt(11)

    doc.add_paragraph()

    # meta table — values derived from DDR data where possible
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.style = "Table Grid"
    cells = [
        ("Inspection Date", "See Additional Notes"),
        ("Inspected By", "See Additional Notes"),
        ("Property Type", "See Additional Notes"),
        ("Areas Inspected", str(len(ddr.get("area_wise_observations", [])))),
        ("Severity Items", str(len(ddr.get("severity_assessment", [])))),
        ("Recommended Actions", str(len(ddr.get("recommended_actions", [])))),
    ]
    for i, (label, val) in enumerate(cells):
        row = i // 2
        col = (i % 2) * 2
        c1 = meta_table.rows[row].cells[col]
        c2 = meta_table.rows[row].cells[col + 1]
        set_cell_bg(c1, "D6E4F7")
        c1.text = label
        c1.paragraphs[0].runs[0].bold = True
        c2.text = val

    doc.add_paragraph()


    add_section_heading(doc, "1. Property Issue Summary")
    doc.add_paragraph(ddr.get("property_issue_summary", "Not Available"))

    doc.add_paragraph()

    # ── Section 2: Area-wise Observations ──
    add_section_heading(doc, "2. Area-wise Observations")

    areas = ddr.get("area_wise_observations", [])
    inspection_page_map = inspection_page_map or {}
    thermal_page_map = thermal_page_map or {}
    insp_img_map, therm_img_map = build_area_image_maps(
        areas,
        inspection_images,
        thermal_images,
        inspection_page_map,
        thermal_page_map,
    )

    for i, obs in enumerate(areas):
        area_name = obs.get("area", f"Area {i+1}")
        add_section_heading(doc, f"2.{i+1}  {area_name}", level=2)

        # Observation table
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        headers = ["Negative Side (Interior)", "Positive Side (Source)", "Thermal Finding", "Image Reference"]
        values = [
            obs.get("negative_side", "Not Available"),
            obs.get("positive_side", "Not Available"),
            obs.get("thermal_finding", "Not Available"),
            obs.get("image_reference", "Image Not Available"),
        ]
        for row_i, (h, v) in enumerate(zip(headers, values)):
            c1 = tbl.rows[row_i].cells[0]
            c2 = tbl.rows[row_i].cells[1]
            set_cell_bg(c1, "EBF3FB")
            c1.text = h
            c1.paragraphs[0].runs[0].bold = True
            c2.text = v

        doc.add_paragraph()

        # Add a thermal image if available
        if i in therm_img_map:
            timg = therm_img_map[i]
            p = doc.add_paragraph()
            run = p.add_run()
            try:
                run.add_picture(timg["path"], width=Inches(3.5))
                cap = doc.add_paragraph(f"Thermal Image — {area_name}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            except Exception as e:
                print(f"Warning: could not embed thermal image {timg['filename']}: {e}")
                p.add_run(f"[Thermal image: {timg['filename']}]")
        else:
            doc.add_paragraph("Thermal Image: Image Not Available")

        # Add inspection photos for this area
        imgs_for_area = insp_img_map.get(i, [])
        if imgs_for_area:
            row_para = doc.add_paragraph()
            for img_info in imgs_for_area[:2]:
                run = row_para.add_run()
                try:
                    run.add_picture(img_info["path"], width=Inches(2.5))
                    row_para.add_run("  ")
                except Exception as e:
                    print(f"Warning: could not embed image {img_info['filename']}: {e}")
                    row_para.add_run(f"[{img_info['filename']}]  ")
            if imgs_for_area:
                cap = doc.add_paragraph(f"Site Photos — {area_name}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph("Inspection Images: Image Not Available")

        doc.add_paragraph()

    # ── Section 3: Probable Root Cause ──
    add_section_heading(doc, "3. Probable Root Cause")
    doc.add_paragraph(ddr.get("probable_root_cause", "Not Available"))
    doc.add_paragraph()

    # ── Section 4: Severity Assessment ──
    add_section_heading(doc, "4. Severity Assessment")
    sev_list = ddr.get("severity_assessment", [])
    if sev_list:
        tbl = doc.add_table(rows=1 + len(sev_list), cols=3)
        tbl.style = "Table Grid"
        headers = ["Area", "Severity", "Reasoning"]
        for j, h in enumerate(headers):
            cell = tbl.rows[0].cells[j]
            set_cell_bg(cell, "1F497D")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row_i, item in enumerate(sev_list, start=1):
            cells = tbl.rows[row_i].cells
            cells[0].text = item.get("area", "")
            sev = item.get("severity", "")
            sev_cell = cells[1]
            p = sev_cell.paragraphs[0]
            run = p.add_run(sev)
            run.bold = True
            color = SEVERITY_COLORS.get(sev, RGBColor(0x66, 0x66, 0x66))
            run.font.color.rgb = color
            cells[2].text = item.get("reasoning", "")

    doc.add_paragraph()

    # ── Section 5: Recommended Actions ──
    add_section_heading(doc, "5. Recommended Actions")
    actions = ddr.get("recommended_actions", [])
    if actions:
        tbl = doc.add_table(rows=1 + len(actions), cols=3)
        tbl.style = "Table Grid"
        for j, h in enumerate(["Priority", "Area", "Action"]):
            cell = tbl.rows[0].cells[j]
            set_cell_bg(cell, "1F497D")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row_i, item in enumerate(actions, start=1):
            cells = tbl.rows[row_i].cells
            pri = item.get("priority", "")
            p = cells[0].paragraphs[0]
            run = p.add_run(pri)
            run.bold = True
            color = PRIORITY_COLORS.get(pri, RGBColor(0x66, 0x66, 0x66))
            run.font.color.rgb = color
            cells[1].text = item.get("area", "")
            cells[2].text = item.get("action", "")

    doc.add_paragraph()

    # ── Section 6: Additional Notes ──
    add_section_heading(doc, "6. Additional Notes")
    doc.add_paragraph(ddr.get("additional_notes", "Not Available"))
    doc.add_paragraph()

    # ── Section 7: Missing / Unclear Information ──
    add_section_heading(doc, "7. Missing or Unclear Information")
    missing = ddr.get("missing_or_unclear_information", ["Not Available"])
    for item in missing:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()

    # ── Footer note ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("— Generated by AI DDR System | Based on provided Inspection & Thermal Reports —")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(out_path)
    print(f"DDR saved to: {out_path}")
    return out_path



# PIPELINE
def generate_ddr(inspection_pdf: str, thermal_pdf: str, api_key: str, out_path: str, allow_fallback: bool = True):
    if not os.path.isfile(inspection_pdf):
        raise FileNotFoundError(f"Inspection PDF not found: {inspection_pdf}")
    if not os.path.isfile(thermal_pdf):
        raise FileNotFoundError(f"Thermal PDF not found: {thermal_pdf}")

    print("Step 1: Extracting text...")
    inspection_text = extract_text(inspection_pdf)
    thermal_text = extract_text(thermal_pdf)
    inspection_page_map = parse_paged_text(inspection_text)
    thermal_page_map = parse_paged_text(thermal_text)
    print(f"  Inspection text: {len(inspection_text)} chars")
    print(f"  Thermal text: {len(thermal_text)} chars")

    print("Step 2: Extracting images...")
    insp_img_dir = os.path.join(IMG_DIR, "inspection")
    therm_img_dir = os.path.join(IMG_DIR, "thermal")
    inspection_images = extract_inspection_images(inspection_pdf, insp_img_dir)
    thermal_images = render_thermal_pages(thermal_pdf, therm_img_dir)
    print(f"  Inspection images: {len(inspection_images)}")
    print(f"  Thermal pages: {len(thermal_images)}")

    print("Step 3: Calling Gemini for DDR analysis...")
    try:
        ddr = call_gemini_for_ddr(inspection_text, thermal_text, api_key)
    except RuntimeError as exc:
        if not allow_fallback:
            raise
        print(f"  Warning: {exc}")
        print("  Switching to fallback DDR generation mode...")
        ddr = build_fallback_ddr(inspection_text, thermal_text, str(exc))
    ddr = normalize_ddr_payload(ddr, inspection_text, thermal_text)
    print(f"  DDR sections: {list(ddr.keys())}")

    # Save raw JSON for debugging
    debug_json_path = str(BASE_DIR / "ddr_raw.json")
    with open(debug_json_path, "w") as f:
        json.dump(ddr, f, indent=2)

    print("Step 4: Building DOCX report...")
    build_docx_report(
        ddr,
        inspection_images,
        thermal_images,
        out_path,
        inspection_page_map=inspection_page_map,
        thermal_page_map=thermal_page_map,
    )

    return out_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate DDR report from inspection and thermal PDFs.")
    parser.add_argument("--api-key", dest="api_key", default=GEMINI_API_KEY, help="Gemini API key")
    parser.add_argument("--inspection", dest="inspection_pdf", default=ENV_INSPECTION_PDF or None, help="Path to inspection PDF")
    parser.add_argument("--thermal", dest="thermal_pdf", default=ENV_THERMAL_PDF or None, help="Path to thermal PDF")
    parser.add_argument(
        "--allow-fallback",
        dest="allow_fallback",
        action=argparse.BooleanOptionalAction,
        default=ENV_ALLOW_FALLBACK,
        help="Allow fallback DDR generation when Gemini API is unavailable",
    )
    parser.add_argument(
        "--out",
        dest="out_path",
        default=ENV_DDR_OUT_PATH or os.path.join(OUT_DIR, "DDR_Report.docx"),
        help="Output .docx path",
    )
    args = parser.parse_args()

    if not args.api_key:
        print("ERROR: Missing Gemini API key. Set GEMINI_API_KEY in .env or pass --api-key.")
        raise SystemExit(1)

    inspection_pdf = args.inspection_pdf
    thermal_pdf = args.thermal_pdf
    if not inspection_pdf or not thermal_pdf:
        discovered_inspection, discovered_thermal = auto_discover_pdfs(BASE_DIR)
        inspection_pdf = inspection_pdf or discovered_inspection
        thermal_pdf = thermal_pdf or discovered_thermal

    if not inspection_pdf or not thermal_pdf:
        print("ERROR: Could not find both inspection and thermal PDFs.")
        print("Provide them explicitly:")
        print("  python ddr.py --inspection <inspection.pdf> --thermal <thermal.pdf>")
        print("Or set INSPECTION_PDF and THERMAL_PDF in .env")
        raise SystemExit(1)

    try:
        generate_ddr(
            inspection_pdf=inspection_pdf,
            thermal_pdf=thermal_pdf,
            api_key=args.api_key,
            out_path=args.out_path,
            allow_fallback=args.allow_fallback,
        )
    except FileNotFoundError as exc:
        print(f"ERROR: {exc}")
        raise SystemExit(1)
    except RuntimeError as exc:
        print(f"ERROR: {exc}")
        raise SystemExit(1)
